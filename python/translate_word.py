#-----------------------------------------------------
#概要：Wordを翻訳するプログラム
#補足：複数パラグラフを翻訳する
#　　　本Pythonファイルと同じフォルダに翻訳元Wordを格納しておく
#　　　翻訳後のファイル名は「翻訳元ファイル名(Translated).拡張子」とする
#-----------------------------------------------------

from docx import Document
import os
from time import sleep
from googletrans import Translator
from openpyxl.descriptors.base import String
 
#-----------------------------------------------------
#翻訳実行メソッド
#-----------------------------------------------------
def translate(trans_obj, file_path: String, src_lang_code: String, dest_lang_code: String):
    trans_obj.updateLog('Start translation.\n')
    document = Document(file_path)
    translator = Translator() #翻訳準備
    total_paragraph_num= len(document.paragraphs) #パラグラフ数
    trans_obj.updateLog('Translating (Number of all paragraphs:' + str(total_paragraph_num) + ')...\n') #翻訳中である旨を表示。

    for i in range(0, total_paragraph_num):
        #10の倍数の時、何パラグラフ目を翻訳中か表示
        if (i+1) % 10 == 0:
            print(str(i + 1) + "/" + str(total_paragraph_num) + "パラグラフ目を翻訳中です...") 
            trans_obj.updateLog('Translating ' + str(i + 1) + "/" + str(total_paragraph_num) + ' paragraph.../n')   #翻訳中である旨を表示。

        str_in=document.paragraphs[i].text #パラグラフのテキスト取得

        if str(str_in)=="":
            continue
        else:
            str_out = translator.translate(str_in, src=src_lang_code, dest=dest_lang_code).text
            document.paragraphs[i].text=str_out
            sleep(1) #サーバ負荷軽減処理。
                
    file_name_split = os.path.splitext(os.path.basename(file_path)) #翻訳元ファイル名を拡張子以外と拡張子に分割
    document.save(file_name_split[0] + "(Translated)" + file_name_split[1])
    trans_obj.updateLog('Translation was completed.\n') #翻訳終了メッセージ
